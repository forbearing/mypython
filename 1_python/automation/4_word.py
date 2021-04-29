---
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn

doc1 = Document()
doc1.add_heading("如何使用 Python 创建 Word", 0)
doc1.save("word1.docx")

https://zhuanlan.zhihu.com/p/61340025
1:新建文档
    from docx import Document
    doc = Document()

2:编辑已经存在的文档
    from docx import Document
    doc = Document('existing-document-file.docx')
    doc.save('new-file-name.docx')

3:新增段落
    paragraph = document.add_paragraph('这是一个段落')
    prior_paragraph = paragraph.insert_paragraph_before('这是前面的段落')
        # 在此段落之前插入一个段落

4:新增标题
    # 新增标题大小，有1-9种规格，如果使用 level=0，则会新增一个带有下划线样式的标题
    doc.add_heading("这是标题")
    doc.add_heading("这是标题", level=2)

5:新增分页符
    doc.add_page_break()

6:新增表格
    table = doc.add_table(rows=2, cols=2)
        # 创建一个2行2列的表格
    cell = table.cell(0,1)
        # 获取第一行和第二行的单元格类
    cell.text = '这是第一行第二列的单元格'
        # 写入数据
    row = table.rows[1]
    row.cells[0].text = '第二行第一列'
    row.cells[1].text = '第二行第二列'
        # 不仅如此，还能以数组的形式获取整个行或列
    for row in table.rows:
        for cell in row.cells:
            print(cell.text)
        # 循环操作
    row_count = len(table.rows)
    col_count = len(table.columns)
        # 用 len() 方法获取行数或列数
    row = table.add_row()
        # 增加行
    table.style = 'LightShading-Accent1'
        # 设置表格样式

7:插入图片
    doc.add_picture('demo.png')             # 插入本地
    from docx.shared import Inches
    doc.add_picture('demo.png', width=Inches(1.0), height=Inches(1.0))
